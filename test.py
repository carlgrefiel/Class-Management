from docx import Document

def convert_docx_to_txt(docx_path, txt_output_path):
    doc = Document(docx_path)
    text_parts = []

    # Extract paragraph text
    for para in doc.paragraphs:
        text_parts.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            text_parts.append("\t".join(row_text))  # Tab-separated for readability

    # Write to .txt file
    with open(txt_output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(text_parts))

    print(f"✅ Converted: {docx_path} → {txt_output_path}")
    
convert_docx_to_txt("example.docx", "output.txt")
