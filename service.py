import os
from docx import Document
import re
import pandas as pd
from dateutil.parser import parse

def extract_birth_line(full_text, validation):
    lines = full_text.splitlines()
    for line in lines:
        if line.strip().startswith(validation):
            return line.strip()
    return None

def format_date(date_str):
    try:
        dt = parse(date_str, dayfirst=False, fuzzy=True)
        return dt.strftime("%#m/%#d/%Y")  # Use "%-m/%-d/%Y" on macOS/Linux
    except Exception:
        return None

def remove_extra_commas(text):
    return re.sub(r',+', ',', text.strip(','))

def smart_split(text):
    cleaned = re.sub(r'[\t]+| {2,}', ',', text.strip())
    return remove_extra_commas(cleaned)

def is_do_variation(s):
    cleaned = s.replace(" ", "").replace("-", "").lower()
    return cleaned in ("do", "d o")

def normalize_status(s):
    s_lower = s.lower()
    if s_lower in ("perm.", "perm"):
        return "Permanent"
    elif s_lower in ("temp.", "temp"):
        return "Temporary"
    elif s_lower in ("sub.", "sub"):
        return "Sub"
    elif s_lower == "casual":
        return "Casual"
    return s.strip()

def extract_info_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text])

    # Extract Name
    name_match = re.search(r'NAME:\s*(.*)', full_text, re.IGNORECASE)
    if name_match:
        full_name = smart_split(name_match.group(1))
        parts = full_name.split(',')
        surname = parts[0].strip() if len(parts) > 0 else None
        given_name = parts[1].strip() if len(parts) > 1 else None
        middle_initial = parts[2].strip() if len(parts) > 2 else None
    else:
        surname = given_name = middle_initial = None
        print(f"⚠️ NAME not found in: {docx_path}")

    # Extract Birth
    birth_match = re.search(r'BIRTH:\s*([A-Za-z]+\s+\d{1,2},\s+\d{4})\s+(.*)', full_text)
    test_birth = extract_birth_line(full_text, "BIRTH:")
    print(test_birth)
    if birth_match:
        raw_birthdate = birth_match.group(1).strip()
        birthplace = birth_match.group(2).strip()
        if birthplace.startswith("Date"):
            birthplace = None
        if birthplace and birthplace.endswith(","):
            birthplace += " Bohol"
        birthdate = format_date(raw_birthdate) or raw_birthdate
    else:
        birthdate = birthplace = None

    # Prepare service record tracking
    service_records = []
    last_valid_status = None
    last_valid_position = None
    last_valid_company = None
    last_valid_branch = None

    for table in doc.tables:
        total_rows = len(table.rows)
        for i, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            if len(row_data) >= 9:
                from_date = format_date(row_data[0])
                if not from_date:
                    print(f"⚠️ Skipping row {i} due to invalid From date: '{row_data[0]}' in {docx_path}")
                    continue

                to_date = format_date(row_data[1]) or row_data[1]

                raw_position = row_data[2]
                raw_status = row_data[3]
                raw_salary = row_data[4]
                raw_company = row_data[5]
                raw_branch = row_data[6]
                lwp = row_data[7]
                separation = row_data[8]

                # Handle status
                if is_do_variation(raw_status):
                    status = last_valid_status
                else:
                    status = normalize_status(raw_status)
                    last_valid_status = status

                # Handle position
                if is_do_variation(raw_position):
                    position = last_valid_position
                else:
                    position = raw_position
                    last_valid_position = position

                # Handle company
                if is_do_variation(raw_company):
                    company = last_valid_company
                else:
                    company = raw_company
                    last_valid_company = company

                # Handle branch
                if is_do_variation(raw_branch):
                    branch = last_valid_branch
                else:
                    branch = raw_branch
                    last_valid_branch = branch

                service_records.append({
                    "First Name": given_name,
                    "Last Name": surname,
                    "Middle Name": middle_initial,
                    "Birthdate": birthdate,
                    "Birthplace": birthplace,
                    "From": from_date,
                    "To": to_date,
                    "Position": position,
                    "Status": status,
                    "Salary": raw_salary,
                    "Company": company,
                    "Branch": branch,
                    "LWP": lwp,
                    "Separation": separation
                })
            else:
                print(f"⚠️ Skipping row {i} with insufficient columns ({len(row_data)}) in {docx_path}")

    return service_records

def process_all_docx_in_folder(folder_path, output_excel_path):
    all_records = []
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(".docx"):
                full_path = os.path.join(root, file)
                print(f"Processing: {full_path}")
                records = extract_info_from_docx(full_path)
                all_records.extend(records)

    if all_records:
        df = pd.DataFrame(all_records)
        df.to_excel(output_excel_path, index=False)
        print(f"\n✅ Combined Excel file saved to: {output_excel_path}")
    else:
        print("\n⚠️ No valid service records found in any document.")

# Example usage
pathFolder = input("Please input file path: ").strip()
folder_path = pathFolder  # Directly use the user input
output_excel_path = r"C:\Users\CodingPh\laravel9\DICT\combined_service_records.xlsx"

# Make sure the function process_all_docx_in_folder is defined somewhere above this
process_all_docx_in_folder(folder_path, output_excel_path)
