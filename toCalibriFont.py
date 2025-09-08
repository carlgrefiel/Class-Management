import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

def set_font_to_calibri(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

def convert_folder_fonts_to_calibri_recursive(input_folder, output_folder):
    for root, _, files in os.walk(input_folder):
        for file in files:
            if file.lower().endswith(".docx"):
                input_file_path = os.path.join(root, file)

                # Mirror subdirectory in output
                relative_path = os.path.relpath(root, input_folder)
                output_dir = os.path.join(output_folder, relative_path)
                os.makedirs(output_dir, exist_ok=True)

                output_file_path = os.path.join(output_dir, file)

                print(f"📄 Converting font in: {input_file_path}")
                try:
                    doc = Document(input_file_path)
                    set_font_to_calibri(doc)
                    doc.save(output_file_path)
                    print(f"✅ Saved: {output_file_path}")
                except Exception as e:
                    print(f"❌ Failed to process {input_file_path}: {e}")

# Example usage
input_folder = r"\\192.168.30.93\carl\Yna\Updateddocs"
output_folder = r"\\192.168.30.93\carl\convertedDoc\calibriFont\test"

convert_folder_fonts_to_calibri_recursive(input_folder, output_folder)
